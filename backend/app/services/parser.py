"""
Document Parser Service
Extracts text from PDF, DOCX, and TXT files
"""

import logging
import os
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles text extraction from different file formats"""

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> Tuple[str, dict]:
        """
        Extract text from a document file.
        Returns (text, metadata) tuple.
        metadata may include page count, author, etc.
        """
        file_type = file_type.lower().strip(".")

        if file_type == "pdf":
            return DocumentParser._parse_pdf(file_path)
        elif file_type == "docx":
            return DocumentParser._parse_docx(file_path)
        elif file_type == "txt":
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _parse_pdf(file_path: str) -> Tuple[str, dict]:
        """Extract text from PDF using PyMuPDF (fitz) - fast and accurate"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            pages_text = []
            for page_num, page in enumerate(doc, 1):
                text = page.get_text("text")
                if text.strip():
                    pages_text.append(f"[Page {page_num}]\n{text}")

            full_text = "\n\n".join(pages_text)
            metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
            }
            doc.close()
            return full_text, metadata
        except ImportError:
            # Fallback to pypdf if fitz not available
            return DocumentParser._parse_pdf_pypdf(file_path)

    @staticmethod
    def _parse_pdf_pypdf(file_path: str) -> Tuple[str, dict]:
        """Fallback PDF parser using pypdf"""
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages_text = []
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(f"[Page {page_num}]\n{text}")

        full_text = "\n\n".join(pages_text)
        metadata = {
            "page_count": len(reader.pages),
            "title": reader.metadata.get("/Title", "") if reader.metadata else "",
            "author": reader.metadata.get("/Author", "") if reader.metadata else "",
        }
        return full_text, metadata

    @staticmethod
    def _parse_docx(file_path: str) -> Tuple[str, dict]:
        """Extract text from DOCX preserving structure"""
        from docx import Document
        doc = Document(file_path)

        sections = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Preserve heading structure
                if para.style.name.startswith("Heading"):
                    sections.append(f"\n## {para.text}\n")
                else:
                    sections.append(para.text)

        # Also extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))
            if table_rows:
                sections.append("\n[Table]\n" + "\n".join(table_rows) + "\n")

        full_text = "\n".join(sections)
        metadata = {
            "page_count": None,  # DOCX doesn't have reliable page count
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
        }
        return full_text, metadata

    @staticmethod
    def _parse_txt(file_path: str) -> Tuple[str, dict]:
        """Read plain text file with encoding detection"""
        # Try UTF-8 first, fall back to latin-1
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                return text, {"page_count": None, "encoding": encoding}
            except UnicodeDecodeError:
                continue

        raise ValueError("Could not decode text file with any common encoding")
